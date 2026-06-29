from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("docs/plan_global_nautilus_ibkr.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
CAUTION_FILL = "FFF2CC"
GREEN_FILL = "E2F0D9"
BORDER = "B8C4D3"


def set_font(run, name: str = "Calibri", size: int | float = 11, bold: bool = False, color: str = "000000") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 60, bottom: int = 60, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            dxa = int(widths_in[idx] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def set_para_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, after=2, line=1.0)
    run = p.add_run("Plan global del proyecto NautilusTrader + IBKR Paper")
    set_font(run, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    set_para_spacing(p, after=10, line=1.15)
    run = p.add_run("Hoja de ruta de trabajo, estado actual y tareas por fase")
    set_font(run, size=12, color=MUTED)

    p = doc.add_paragraph()
    set_para_spacing(p, after=12, line=1.15)
    run = p.add_run(f"Fecha: {date.today().strftime('%d/%m/%Y')}  |  Entorno: TWS paper, sin Docker")
    set_font(run, size=9.5, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix) :])
        set_font(r)
    else:
        set_font(p.add_run(text))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_spacing(p, after=4)
    set_font(p.add_run(text))


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_spacing(p, after=4)
    set_font(p.add_run(text))


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=2, line=1.2)
    r = p.add_run(label)
    set_font(r, bold=True, color=INK)
    r = p.add_run(" " + text)
    set_font(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0, line=1.1)
        r = p.add_run(header)
        set_font(r, size=9.5, bold=True, color=INK)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_borders(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_para_spacing(p, after=0, line=1.15)
            r = p.add_run(value)
            set_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Plan NautilusTrader + IBKR Paper | Hoja de ruta")
    set_font(run, size=8.5, color=MUTED)


def build_doc() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    add_callout(
        doc,
        "Objetivo global:",
        "construir una base de trading paper con NautilusTrader e Interactive Brokers que permita trabajar a dos personas, probar estrategias, registrar resultados y reducir riesgos antes de pensar en real.",
        GREEN_FILL,
    )

    add_heading(doc, "1. Estado actual", 1)
    add_table(
        doc,
        ["Area", "Estado", "Que significa"],
        [
            ["Entorno local", "Hecho", "Proyecto con entorno Python, uv y NautilusTrader instalado con soporte IBKR."],
            ["Docker", "Descartado", "El flujo queda basado en TWS/IB Gateway local, sin contenedores."],
            ["TWS paper", "Hecho", "API socket activada en puerto paper 7497 y cuenta demo validada."],
            ["Conexion IBKR", "Hecho", "El script de conexion ya ha devuelto cuenta, posiciones y ordenes abiertas."],
            ["Primera orden paper", "Hecho", "Se ejecuto una compra paper de 1 AAPL como prueba controlada."],
            ["Documentacion", "En curso", "Hay README, guia de colaboracion y manual basico de TWS."],
            ["Estrategias", "Pendiente", "Todavia no hay estrategia NautilusTrader automatica ni backtest."],
            ["Logs y control", "Pendiente", "Aun falta registrar ordenes, ejecuciones, errores y resultados de forma limpia."],
        ],
        [1.35, 1.15, 4.0],
    )

    add_heading(doc, "2. Orden global recomendado", 1)
    for item in [
        "Cerrar la base del proyecto: estructura, README, .env.example y Git.",
        "Anadir seguridad paper-only y limites basicos antes de nuevas ordenes.",
        "Crear logs de ordenes, ejecuciones, conexiones y errores.",
        "Mejorar las herramientas manuales: comprar, vender, revisar posicion, revisar ordenes y consultar datos.",
        "Crear una primera estrategia minima en NautilusTrader.",
        "Preparar un backtest sencillo antes de operar una estrategia en paper.",
        "Ejecutar paper trading controlado con tamanos pequenos y revision posterior.",
        "Formalizar el trabajo de dos personas con ramas, revisiones y configuracion local separada.",
    ]:
        add_number(doc, item)

    add_callout(
        doc,
        "Regla de trabajo:",
        "no saltar a estrategias complejas hasta tener seguridad, logs y comprobaciones. La prioridad es que el sistema sea entendible, repetible y revisable.",
        CAUTION_FILL,
    )

    add_heading(doc, "3. Detalle por punto", 1)

    phases = [
        {
            "title": "Punto 1 - Base del proyecto y Git",
            "objective": "Dejar el proyecto ordenado para que dos personas puedan trabajar sin pisarse configuraciones.",
            "tasks": [
                "Confirmar carpetas: scripts, strategies, config, logs, data y docs.",
                "Mantener .env privado y .env.example como plantilla compartible.",
                "Inicializar Git si aun no esta inicializado.",
                "Crear rama estable main y ramas de trabajo por persona.",
                "Actualizar README con comandos basicos y flujo de instalacion.",
            ],
            "deliverable": "Repositorio limpio, documentado y preparado para compartir codigo sin compartir credenciales.",
            "success": "Un segundo ordenador puede clonar el repo, crear su .env y ejecutar los checks basicos.",
        },
        {
            "title": "Punto 2 - Seguridad paper-only",
            "objective": "Evitar errores antes de enviar mas ordenes desde codigo.",
            "tasks": [
                "Bloquear cualquier cuenta que no empiece por DU salvo permiso explicito.",
                "Exigir flags de confirmacion para transmitir ordenes.",
                "Poner limite maximo por orden y limite de simbolos permitidos para pruebas.",
                "Dejar whatIf como modo por defecto.",
                "Mostrar resumen claro antes de enviar: cuenta, simbolo, lado, cantidad y tipo.",
            ],
            "deliverable": "Capa de seguridad reutilizable por todos los scripts.",
            "success": "Una orden no se transmite si falta confirmacion, si la cuenta no es paper o si supera limites.",
        },
        {
            "title": "Punto 3 - Logs y trazabilidad",
            "objective": "Saber exactamente que se pidio, que acepto IBKR y que se ejecuto.",
            "tasks": [
                "Crear logs/orders.csv para solicitudes de orden.",
                "Crear logs/executions.csv para ejecuciones recibidas.",
                "Crear logs/connection_checks.log para diagnostico.",
                "Guardar timestamp, cuenta, client_id, simbolo, lado, cantidad, tipo, estado y mensaje de error.",
                "Asegurar que logs locales no se suben a Git.",
            ],
            "deliverable": "Registro basico de operaciones y revisiones.",
            "success": "Despues de una prueba se puede reconstruir que paso sin depender solo de TWS.",
        },
        {
            "title": "Punto 4 - Herramientas manuales",
            "objective": "Tener comandos simples para operar y revisar sin tocar el codigo cada vez.",
            "tasks": [
                "Crear comando de compra/venta con parametros claros.",
                "Crear comando para ver posiciones actuales.",
                "Crear comando para ver ordenes abiertas.",
                "Crear comando para consultar ultimo precio, bid/ask y datos basicos.",
                "Anadir mensajes de salida entendibles para una persona no tecnica.",
            ],
            "deliverable": "Kit de comandos para pruebas paper manuales.",
            "success": "Puedes comprar, vender, comprobar posicion y revisar ordenes desde terminal con comandos consistentes.",
        },
        {
            "title": "Punto 5 - Primera estrategia minima",
            "objective": "Validar el flujo de estrategia sin buscar rentabilidad todavia.",
            "tasks": [
                "Crear una estrategia simple y controlada.",
                "Definir instrumento inicial, por ejemplo AAPL o IBKR.",
                "Evitar duplicar compras si ya existe posicion.",
                "Registrar decisiones: por que entra, por que no entra y que orden envia.",
                "Ejecutarla primero en modo simulacion o paper muy pequeno.",
            ],
            "deliverable": "Primera estrategia NautilusTrader ejecutable.",
            "success": "La estrategia toma una decision, envia o no envia orden y deja log de todo.",
        },
        {
            "title": "Punto 6 - Backtest",
            "objective": "Probar ideas con datos historicos antes de dejarlas operar en paper.",
            "tasks": [
                "Definir fuente de datos historicos.",
                "Crear un backtest minimo con comisiones y tamano pequeno.",
                "Guardar resultados: operaciones, P&L, drawdown y numero de trades.",
                "Comparar la logica esperada contra el resultado.",
                "Descartar estrategias que no pasen una prueba minima.",
            ],
            "deliverable": "Pipeline de backtest basico.",
            "success": "Antes de paper trading hay un resultado historico revisable.",
        },
        {
            "title": "Punto 7 - Paper trading controlado",
            "objective": "Ejecutar la estrategia en paper con riesgo limitado y revision diaria.",
            "tasks": [
                "Empezar con 1 accion o tamano minimo.",
                "Ejecutar solo en horarios definidos.",
                "Revisar TWS, logs y posiciones despues de cada prueba.",
                "Cerrar posiciones si la prueba termina o algo queda abierto sin querer.",
                "Documentar resultado y cambios necesarios.",
            ],
            "deliverable": "Rutina repetible de prueba en paper.",
            "success": "Cada ejecucion puede revisarse y repetirse sin dudas sobre que paso.",
        },
        {
            "title": "Punto 8 - Trabajo de dos personas",
            "objective": "Permitir que cada persona desarrolle desde su ordenador, rama y datos propios.",
            "tasks": [
                "Crear ramas separadas por persona o tarea.",
                "No compartir .env, credenciales ni logs privados.",
                "Usar client_id diferente si se conecta al mismo TWS.",
                "Fusionar cambios solo despues de revisar.",
                "Mantener main como version estable.",
            ],
            "deliverable": "Flujo de colaboracion estable.",
            "success": "Cada uno puede probar con su configuracion sin afectar al otro.",
        },
    ]

    for phase in phases:
        if phase["title"].startswith("Punto 8"):
            doc.add_page_break()
        add_heading(doc, phase["title"], 2)
        add_body(doc, "Objetivo: " + phase["objective"], bold_prefix="Objetivo:")
        add_heading(doc, "Tareas", 3)
        for task in phase["tasks"]:
            add_bullet(doc, task)
        add_table(
            doc,
            ["Resultado esperado", "Criterio para darlo por terminado"],
            [[phase["deliverable"], phase["success"]]],
            [3.0, 3.5],
        )

    add_heading(doc, "4. Lo que no haria todavia", 1)
    for item in [
        "No pasaria a cuenta real.",
        "No ejecutaria estrategias automaticas sin logs.",
        "No abriria muchos simbolos a la vez.",
        "No mezclaria pruebas de dos personas sobre el mismo client_id.",
        "No subiria .env, logs privados ni capturas con cuenta visible.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Primer bloque de trabajo inmediato", 1)
    add_table(
        doc,
        ["Orden", "Trabajo", "Resultado"],
        [
            ["1", "Revisar estructura y Git", "Proyecto base preparado para colaborar."],
            ["2", "Seguridad paper-only", "Ordenes bloqueadas si no cumplen reglas."],
            ["3", "Logs", "Cada prueba deja rastro en archivos locales."],
            ["4", "Comandos manuales", "Comprar, vender y revisar estado desde terminal."],
            ["5", "Estrategia minima", "Primer flujo automatico pequeno y controlado."],
        ],
        [0.7, 2.6, 3.2],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    build_doc()
