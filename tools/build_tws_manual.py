from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from PIL import Image
from PIL import ImageDraw
from PIL import ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "assets/screenshots/tws_full_after_activate.png"
OUT_DIR = ROOT / "assets/screenshots/manual"
DOCX_PATH = ROOT / "docs/manual_basico_tws.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3E1"
RISK_RED = "9B1C1C"
GREEN = "1B5E20"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            pass
    return ImageFont.load_default()


def redact_account(img: Image.Image, offset: tuple[int, int] = (0, 0)) -> None:
    draw = ImageDraw.Draw(img)
    ox, oy = offset
    # Original-screen coordinates where the paper account appears in the captured TWS window.
    redactions = [
        (1395, 100, 1545, 130),  # top status/account area
        (675, 955, 790, 985),  # activity/trades account column
    ]
    for x1, y1, x2, y2 in redactions:
        box = (x1 - ox, y1 - oy, x2 - ox, y2 - oy)
        if box[2] > 0 and box[3] > 0 and box[0] < img.width and box[1] < img.height:
            draw.rounded_rectangle(box, radius=4, fill=(18, 18, 18))
            draw.text((box[0] + 6, box[1] + 4), "PAPER", fill=(235, 235, 235), font=font(16, True))


def crop(name: str, box: tuple[int, int, int, int]) -> Image.Image:
    src = Image.open(SOURCE).convert("RGB")
    cropped = src.crop(box)
    redact_account(cropped, offset=(box[0], box[1]))
    path = OUT_DIR / f"{name}.png"
    cropped.save(path)
    return cropped


def label(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, color=(255, 214, 10)) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=8, outline=color, width=5)
    text_font = font(22, True)
    bbox = draw.textbbox((0, 0), text, font=text_font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    pad = 8
    lx = x1
    ly = max(0, y1 - th - pad * 2)
    draw.rounded_rectangle((lx, ly, lx + tw + pad * 2, ly + th + pad * 2), radius=6, fill=(20, 20, 20))
    draw.text((lx + pad, ly + pad), text, fill=(255, 255, 255), font=text_font)


def annotate(name: str, img: Image.Image, labels: list[tuple[tuple[int, int, int, int], str]]) -> Path:
    annotated = img.copy()
    draw = ImageDraw.Draw(annotated)
    for box, text in labels:
        label(draw, box, text)
    path = OUT_DIR / f"{name}_annotated.png"
    annotated.save(path)
    return path


def prepare_images() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    full = crop("01_mosaic_full", (0, 45, 1775, 1165))
    order = crop("02_order_entry", (0, 190, 910, 410))
    portfolio = crop("03_portfolio", (910, 190, 1775, 865))
    market_chart = crop("04_market_chart", (0, 410, 910, 850))
    activity = crop("05_activity_trades", (0, 860, 910, 1120))

    return {
        "full": annotate(
            "01_mosaic_full",
            full,
            [
                ((5, 85, 1760, 145), "Cuenta paper y estado"),
                ((0, 155, 910, 360), "Entrada de ordenes"),
                ((0, 365, 315, 805), "Datos de mercado"),
                ((315, 365, 905, 805), "Grafico"),
                ((910, 155, 1765, 805), "Portfolio"),
                ((0, 805, 910, 1075), "Actividad"),
                ((910, 805, 1765, 1075), "Noticias"),
            ],
        ),
        "order": annotate(
            "02_order_entry",
            order,
            [
                ((10, 45, 160, 82), "Ticker"),
                ((15, 150, 180, 196), "BUY / SELL"),
                ((205, 150, 365, 196), "Cantidad"),
                ((365, 150, 675, 196), "Tipo / precio / vigencia"),
                ((790, 145, 900, 198), "Submit"),
                ((200, 50, 300, 125), "Posicion"),
            ],
        ),
        "portfolio": annotate(
            "03_portfolio",
            portfolio,
            [
                ((0, 0, 650, 70), "Monitor / Portfolio"),
                ((35, 100, 300, 225), "P&L"),
                ((35, 225, 720, 325), "Margen y liquidez"),
                ((35, 330, 725, 440), "Posiciones"),
            ],
        ),
        "market_chart": annotate(
            "04_market_chart",
            market_chart,
            [
                ((0, 0, 315, 435), "Quote panel"),
                ((320, 0, 905, 435), "Grafico"),
                ((0, 95, 315, 250), "Bid / Ask / Last"),
                ((0, 280, 315, 390), "Rangos y datos"),
            ],
        ),
        "activity": annotate(
            "05_activity_trades",
            activity,
            [
                ((0, 0, 520, 55), "Orders / Trades / Summary"),
                ((45, 65, 850, 135), "Ejecuciones"),
                ((0, 135, 900, 245), "Zona de seguimiento"),
            ],
        ),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_run_font(run, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_start: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        run = p.add_run(bold_start)
        run.bold = True
        set_run_font(run)
        rest = p.add_run(text[len(bold_start) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, BORDER)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_run_font(r)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_run_font(r)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_width(table, [1.8, 4.5])
    header = table.rows[0].cells
    header[0].text = "Elemento"
    header[1].text = "Uso basico"
    for cell in header:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_run_font(r)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Manual basico TWS | Cuenta paper | ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)
    set_run_font(run)


def build_doc(images: dict[str, Path]) -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Manual basico de Trader Workstation")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_run_font(r)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Uso inicial para cuenta paper: interfaz, datos, compra/venta y seguimiento")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)
    set_run_font(r)

    add_callout(
        doc,
        "Alcance",
        "Guia practica y breve. No cubre toda TWS; se centra en lo necesario para operar en paper, revisar datos y evitar errores basicos.",
    )

    doc.add_heading("1. Antes de operar", level=1)
    add_bullets(
        doc,
        [
            "Confirma que arriba aparece la banda roja de paper trading. Si no aparece, no sigas.",
            "Verifica que la cuenta sea paper/demo y que el simbolo seleccionado sea el correcto.",
            "Empieza con ordenes pequenas. La cuenta paper simula ejecuciones, pero el flujo es el mismo que en real.",
            "Si vas a usar NautilusTrader, TWS debe estar abierto y la API activada en File > Global Configuration > API > Settings.",
        ],
    )

    add_key_value_table(
        doc,
        [
            ("Puerto API", "TWS paper suele usar 7497. IB Gateway paper suele usar 4002."),
            ("Read-Only API", "Activado para pruebas de conexion. Desactivado solo cuando quieras permitir ordenes desde la API."),
            ("Orden de prueba", "BUY 1 AAPL MKT paper fue el primer flujo validado en este entorno."),
            ("Revision posterior", "Comprobar Portfolio, Activity > Trades y open orders."),
        ],
    )

    doc.add_heading("2. Vista principal Mosaic", level=1)
    add_para(
        doc,
        "Mosaic agrupa los bloques basicos en una sola pantalla: entrada de ordenes, datos del instrumento, grafico, portfolio, actividad y noticias.",
    )
    add_image(doc, images["full"], "Vista Mosaic anotada: zonas principales de trabajo.", width=6.4)

    doc.add_heading("3. Entrada de ordenes: compra y venta", level=1)
    add_image(doc, images["order"], "Order Entry: controles minimos antes de pulsar Submit.", width=6.4)
    add_numbered(
        doc,
        [
            "Escribe o confirma el ticker en la caja superior. Ejemplo: AAPL.",
            "Elige BUY o SELL. BUY abre/aumenta una posicion larga; SELL vende o reduce posicion.",
            "Introduce cantidad. Para la primera prueba usamos 1 accion.",
            "Elige tipo de orden: MKT ejecuta al mercado; LMT exige precio limite.",
            "Revisa vigencia: DAY suele ser suficiente para pruebas basicas.",
            "Pulsa Submit solo cuando ticker, lado, cantidad, tipo y cuenta sean correctos.",
        ],
    )
    add_callout(
        doc,
        "Regla practica",
        "Para empezar, usa MKT solo en paper o con tamanos pequenos. Para real, normalmente preferiras LMT para controlar el precio maximo de compra o minimo de venta.",
        fill="FFF2CC",
    )

    doc.add_heading("4. Portfolio: posicion, P&L y margen", level=1)
    add_image(doc, images["portfolio"], "Portfolio: posicion actual, P&L, liquidez y margen.", width=6.4)
    add_key_value_table(
        doc,
        [
            ("P&L diario", "Resultado del dia. Puede moverse aunque no cierres la posicion."),
            ("Unrealized", "Ganancia/perdida latente de posiciones abiertas."),
            ("Realized", "Ganancia/perdida ya cerrada."),
            ("Net Liquidity", "Valor total aproximado de la cuenta."),
            ("Pos", "Cantidad actual por instrumento. Ejemplo: 1 AAPL."),
            ("Avg Px", "Precio medio de la posicion."),
        ],
    )

    doc.add_heading("5. Datos de mercado y grafico", level=1)
    add_image(doc, images["market_chart"], "Quote panel y grafico: precio, bid/ask, rangos y velas.", width=6.4)
    add_bullets(
        doc,
        [
            "Last: ultimo precio negociado.",
            "Bid/Ask: mejor precio comprador/vendedor disponible.",
            "Bid/Ask Size: tamano disponible en cada lado.",
            "Hi/Lo y 52 H/L: rangos recientes utiles para contexto.",
            "Grafico: ayuda a ver tendencia, volatilidad y niveles, pero no sustituye una regla de trading.",
        ],
    )

    doc.add_heading("6. Activity: ordenes y ejecuciones", level=1)
    add_image(doc, images["activity"], "Activity: pestanas de Orders, Trades y Summary.", width=6.4)
    add_bullets(
        doc,
        [
            "Orders: ordenes abiertas, pendientes, canceladas o rechazadas.",
            "Trades: ejecuciones realizadas. Tras la compra paper, aqui aparece AAPL, cantidad 1 y precio de ejecucion.",
            "Summary: resumen operativo de actividad.",
            "Si no ves ejecucion, revisa si la orden quedo abierta, fue rechazada o requiere confirmacion en TWS.",
        ],
    )

    doc.add_heading("7. Flujo basico recomendado", level=1)
    add_numbered(
        doc,
        [
            "Abrir TWS en paper y confirmar banda roja de simulated trading.",
            "Confirmar API si se va a usar NautilusTrader.",
            "Buscar ticker y revisar datos basicos: Last, Bid, Ask, volumen/tamano y grafico.",
            "Preparar la orden en paper: lado, cantidad, tipo, precio si aplica y vigencia.",
            "Enviar orden y revisar Activity > Trades.",
            "Revisar Portfolio para confirmar posicion y P&L.",
            "Cerrar o reducir posicion solo con una orden contraria clara.",
        ],
    )

    doc.add_heading("8. Errores comunes", level=1)
    add_key_value_table(
        doc,
        [
            ("Ticker equivocado", "Comprueba simbolo y mercado antes de Submit."),
            ("Cantidad incorrecta", "Revisar cantidad antes de enviar. En paper tambien conviene disciplina."),
            ("MKT sin mirar spread", "Si bid/ask esta muy abierto, usa limite o espera mejor liquidez."),
            ("Confundir paper con real", "La banda roja de paper debe estar visible durante pruebas."),
            ("No revisar ejecucion", "Despues de enviar, mira Trades y Portfolio. No asumas que se ejecuto."),
        ],
    )

    add_callout(
        doc,
        "Siguiente paso para el proyecto",
        "Con la conexion y una orden paper ya validadas, el siguiente bloque de trabajo es crear una estrategia simple en NautilusTrader, ejecutar backtest o paper-run controlado y registrar resultados.",
        fill=LIGHT_GRAY,
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    images = prepare_images()
    build_doc(images)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
